"""Helpers for building minimal in-memory PDF/DOCX files for tests, so we
don't need to check binary sample files into the repo."""

import io

from docx import Document
from pypdf import PdfWriter


def build_pdf_with_text(text: str) -> bytes:
    content = f"BT /F1 24 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >>\nendobj\n",
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        f"5 0 obj\n<< /Length {len(content)} >>\nstream\n".encode("latin-1")
        + content
        + b"\nendstream\nendobj\n",
    ]

    pdf = b"%PDF-1.4\n"
    offsets = []
    for obj in objects:
        offsets.append(len(pdf))
        pdf += obj
    xref_offset = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n".encode()
    pdf += b"0000000000 65535 f \n"
    for offset in offsets:
        pdf += f"{offset:010d} 00000 n \n".encode()
    pdf += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF"
    ).encode()
    return pdf


def build_image_only_pdf() -> bytes:
    buf = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(buf)
    return buf.getvalue()


def build_docx_with_text(text: str) -> bytes:
    buf = io.BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buf)
    return buf.getvalue()


def build_empty_docx() -> bytes:
    buf = io.BytesIO()
    Document().save(buf)
    return buf.getvalue()
